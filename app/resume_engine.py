from docx import Document
from pathlib import Path
from ats_scoring import extract_keywords, ats_score

BASE_RESUME = "/app/templates/base_resume1.docx"
JD_FILE = "/app/input/jd_links.txt"
OUTPUT_DIR = Path("/app/output")
OUTPUT_DIR.mkdir(exist_ok=True)

# Read JD
jd_text = Path(JD_FILE).read_text()

keywords = extract_keywords(jd_text)
score = ats_score(keywords)

doc = Document(BASE_RESUME)

doc.add_paragraph("\nATS OPTIMIZATION SUMMARY")
doc.add_paragraph(f"Matched Keywords: {', '.join(keywords)}")
doc.add_paragraph(f"Estimated ATS Score: {score}%")

output_file = OUTPUT_DIR / "resume_updated.docx"
doc.save(output_file)

print(f"ATS Score: {score}%")
print("Resume updated successfully")

